from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(r"C:\品味识别\intelifar-ip-wiki-graph")
OUT = ROOT / "deliverables" / "intelifar-企业文档IP智能平台-产品介绍.docx"
ASSETS = ROOT / "artifacts"
BRAND = ROOT / "site" / "public" / "brand"

PURPLE = RGBColor(0x65, 0x57, 0xFF)
PURPLE_DARK = RGBColor(0x45, 0x38, 0xD6)
INK = RGBColor(0x20, 0x22, 0x34)
MUTED = RGBColor(0x6C, 0x70, 0x82)
DEEP = RGBColor(0x17, 0x18, 0x27)
TEAL = RGBColor(0x18, 0xA8, 0x89)
GOLD = RGBColor(0xB2, 0x78, 0x12)
LIGHT = "F5F4FF"
LIGHT_TEAL = "EAF8F4"
LIGHT_GOLD = "FFF6DF"
BORDER = "D9DCE8"
FONT_CN = "Microsoft YaHei"
FONT_EN = "Arial"


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False, font=FONT_CN):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_EN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, after=8, before=0, line=1.333,
             keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_mixed_para(doc, parts, after=8, line=1.333, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, options in parts:
        set_run(p.add_run(text), **options)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level],
            color=PURPLE if level < 3 else PURPLE_DARK, bold=True)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run(p.add_run(text), size=10.6)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.208
    set_run(p.add_run(text), size=10.6)
    return p


def add_kicker(doc, text):
    p = add_para(doc, text.upper(), size=9.5, color=PURPLE, bold=True, after=6,
                 keep=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, title, body, fill=LIGHT, accent=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color=fill, size="0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(title), size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.25
    set_run(p2.add_run(body), size=10.4, color=INK)
    add_para(doc, "", size=2, after=4)
    return table


def add_picture(doc, rel_path, caption, width=6.25):
    path = ROOT / rel_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    doc_props = run._element.xpath(".//wp:docPr")
    if doc_props:
        doc_props[0].set("descr", caption)
    cap = add_para(doc, caption, size=8.8, color=MUTED, italic=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.15)
    return cap


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run(run, size=8.5, color=MUTED)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size, before, after in ((1, 16, 18, 10), (2, 13, 12, 6), (3, 12, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = PURPLE if level < 3 else PURPLE_DARK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("intelifar  |  企业文档 IP 智能平台"), size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    set_run(fp.add_run("对外产品介绍 · 2026"), size=8.5, color=MUTED)
    add_page_number(fp)


def add_two_col_table(doc, rows, widths=(3000, 6360), header=None):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header:
        cells = table.add_row().cells
        for idx, text in enumerate(header):
            shade_cell(cells[idx], "ECEAFF")
            p = cells[idx].paragraphs[0]
            set_run(p.add_run(text), size=10, color=PURPLE_DARK, bold=True)
    for left, right in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((left, right)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.2
            set_run(p.add_run(text), size=9.8, color=INK, bold=(idx == 0))
    set_table_geometry(table, list(widths))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def section_page(doc, kicker, title, intro=None):
    doc.add_page_break()
    add_kicker(doc, kicker)
    add_heading(doc, title, 1)
    if intro:
        add_para(doc, intro, size=11.2, color=MUTED, after=12)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    # Cover - editorial_cover pattern
    add_para(doc, "", size=2, after=64)
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = p_logo.add_run()
    logo_run.add_picture(str(BRAND / "intelifar-logo.png"), width=Inches(2.25))
    logo_props = logo_run._element.xpath(".//wp:docPr")
    if logo_props:
        logo_props[0].set("descr", "intelifar 品牌标志")
    add_para(doc, "企业知识资产化产品介绍", size=10, color=PURPLE, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=34, after=18)
    add_para(doc, "让每份长文档，\n成为可追溯、可治理、可复用的企业资产", size=27,
             color=DEEP, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20, line=1.15)
    add_para(doc, "intelifar 企业文档 IP 智能分析与 Wiki 管理平台",
             size=14, color=PURPLE_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_para(doc, "把解析、结构化、原文证据、知识网络、受控任务代理和企业治理\n连接成一条可以落地的业务链路",
             size=11.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=70, line=1.35)
    add_para(doc, "面向潜在客户与投资人  |  2026 年 8 月",
             size=9.5, color=MUTED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "产品演示与合作洽谈材料",
             size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Executive summary
    section_page(doc, "EXECUTIVE VIEW", "一个产品，解决三类企业知识问题")
    add_callout(doc, "核心价值主张",
                "intelifar 不是给长文档增加一个聊天框，而是把授权文档转化为有来源、有关系、有版本、有权限的企业 IP 知识资产。")
    add_two_col_table(doc, [
        ("找得到", "把 PDF、Word、PPT、图片和 HTML 中的关键知识转化为可搜索资产与 Wiki。"),
        ("信得过", "每项重要结论保留原文位置、摘录和校验信息，支持回到来源复核。"),
        ("用得起来", "通过 IP 全景网络、跨 Wiki 搜索和任务助手完成盘点、影响分析、证据核查与尽调准备。"),
        ("管得住", "角色权限、人工审批、版本记录、脱敏分享、操作留痕和备份构成治理闭环。"),
    ], header=("客户需要的结果", "intelifar 的交付方式"))
    add_heading(doc, "适合率先使用的团队", 2)
    for item in (
        "研发、知识产权和技术管理团队：沉淀技术方案、专利线索、依赖关系和来源依据。",
        "咨询、专业服务和研究团队：把大量项目资料转化为可复用的方法、案例和交付知识。",
        "制造、工程和产品团队：连接标准、方案、产品能力、供应链资料和风险线索。",
        "小微企业负责人：用较低的治理复杂度建立可持续积累的企业知识空间。",
    ):
        add_bullet(doc, item)

    # Problem / workflow
    section_page(doc, "THE PROBLEM", "企业真正缺少的不是更多摘要，而是一条可信的知识链",
                 "文档越多，知识越容易被锁在文件夹、个人电脑和聊天记录中；即使可以搜索，也很难判断结论来自哪里、是否仍有效、谁有权使用。")
    add_two_col_table(doc, [
        ("知识散落", "同一技术、客户或产品信息分布在多种格式与多个版本中。"),
        ("证据断裂", "摘要和回答无法稳定回到原文位置，重要结论难以复核。"),
        ("关系不可见", "技术依赖、复用路径、冲突与影响范围只能靠少数专家记忆。"),
        ("治理滞后", "权限、审批、对外分享和离职交接常在知识生成之后补做。"),
    ], header=("典型障碍", "业务后果"))
    add_heading(doc, "intelifar 的闭环工作方式", 2)
    for step in (
        "接入授权文档：校验格式、大小与风险，保留任务记录。",
        "解析内容结构：识别章节、表格和版面，形成可分析文本。",
        "提取 IP 资产：按业务结构生成资产、风险线索、关系和 Wiki 草稿。",
        "人工复核发布：核对标题、摘要、原文依据、权属与密级。",
        "形成知识网络：连接资产、Wiki、来源、版本和关系。",
        "持续复用治理：搜索、任务助手、安全分享、语义体检与审计。",
    ):
        add_number(doc, step)

    # Product UI
    section_page(doc, "REAL PRODUCT", "从第一份文档开始，建立可以持续增长的 IP 资产库",
                 "指挥台汇总当前工作空间的真实资产、依据覆盖、待办和风险状态，让负责人先看到业务结果，再进入具体治理动作。")
    add_picture(doc, "artifacts/real-ui-port/01-real-ui-home-4388.png",
                "真实工作空间指挥台：资产、来源覆盖、待办和风险状态统一呈现", width=6.25)
    add_mixed_para(doc, [
        ("产品不是单点工具。", {"size": 10.8, "color": PURPLE_DARK, "bold": True}),
        (" 文档中心、智能分析、IP 资产库、IP Wiki、任务助手、风险线索、生命周期、操作记录、成员与权限和系统状态共享同一工作空间与权限体系。",
         {"size": 10.8, "color": INK}),
    ])

    # Evidence Wiki
    section_page(doc, "TRUST BY DESIGN", "每个重要结论，都能沿证据链回到原文",
                 "真正可用于企业决策的知识，不仅要回答“是什么”，还要回答“来自哪里、处于哪个版本、是否经过复核”。")
    add_picture(doc, "artifacts/screenshots/14-real-published-wiki.png",
                "已发布 Wiki 与原文资源并列核对：保留文档、位置、校验信息和引用内容", width=5.4)
    add_two_col_table(doc, [
        ("资产", "承载标题、类型、摘要、权属、密级、置信度和状态。"),
        ("证据", "记录来源文档、章节或页码、摘录与内容校验信息。"),
        ("Wiki", "把分散资产组织成连续可读、可搜索、可版本化的知识页面。"),
        ("版本", "更新只增不改；审批前正式版本保持不变。"),
    ], header=("知识对象", "为什么可信"))

    # Panorama
    section_page(doc, "IP PANORAMA", "知识不再只是页面，而是一张可以缩放和追问的关系网络",
                 "全景网络帮助负责人发现核心资产、技术依赖、复用路径、潜在冲突和变更影响；显示范围始终受当前账号权限约束。")
    add_picture(doc, "artifacts/ip-asset-graph/06-neural-panorama.png",
                "可缩放 IP 资产全景：聚焦核心节点、已确认关系和待复核建议", width=6.15)
    for item in (
        "缩小看全局：识别资产密集区、孤立知识和关键连接点。",
        "放大看细节：查看权属、密级、直接关系、证据数量和来源。",
        "聚焦一跳关系：评估某项技术、规则或方案变化可能影响哪些成果。",
        "人工确认关系：系统建议只有经业务人员复核后才进入正式知识网络。",
    ):
        add_bullet(doc, item)

    # Agent
    section_page(doc, "BOUNDED AGENT", "自然语言负责泛化，确定性边界负责企业交付",
                 "IP 任务助手面向资产盘点、影响分析、证据核查、资产对比、风险缺口、尽调材料和 Wiki 更新建议，不是通用 coding agent。")
    add_picture(doc, "artifacts/ip-agent/02-grounded-delivery.png",
                "受控 IP 任务交付：展示计划、步骤收据、来源入口、待核实项和安全边界", width=5.35)
    add_two_col_table(doc, [
        ("可以做", "搜索授权资产、检查关系、核对原文依据、比较资产、形成结构化分析和草案建议。"),
        ("不会做", "运行代码或系统命令、任意联网、删除或发布知识、修改权限、自动确认关系。"),
        ("如何确定", "每一步使用封闭工具清单和角色权限；结果经过证据门槛并标出待复核项。"),
    ], header=("能力边界", "产品行为"))

    # Governance
    section_page(doc, "ENTERPRISE GOVERNANCE", "知识资产化的最后一公里，是权限、审批和留痕",
                 "intelifar 将治理动作放在业务流程中，而不是在发生问题后补充记录。")
    add_two_col_table(doc, [
        ("角色与权限", "空间所有者、管理员、知识编辑者、只读成员和外部收件人各有清晰边界。"),
        ("人工复核", "发布、关系确认、权属密级、语义重复与冲突建议均保留人工决定。"),
        ("安全分享", "脱敏 Wiki 使用分享链接与独立访问码，并支持有效期与撤销。"),
        ("操作记录", "重要业务动作记录操作人、对象、时间与结果，并支持完整性校验。"),
        ("连续性", "任务、Wiki 版本和工作空间数据持久保存，并提供已验证备份。"),
    ], header=("治理能力", "客户获得什么"))
    add_picture(doc, "artifacts/user-guide-review/06-operation-records.png",
                "操作记录：以业务语言查看重要动作，并保留可核验的记录完整性", width=6.15)

    # Semantic review
    section_page(doc, "KNOWLEDGE QUALITY", "知识库会增长，也需要持续发现重复、冲突和治理欠账",
                 "本地语义资产体检对已发布资产的受限摘要进行检查，把疑似重复和信息不一致转化为管理员待办。")
    add_picture(doc, "artifacts/semantica-phase-2/07-real-action-center.png",
                "真实待办中心：语义资产建议与权属密级治理进入同一工作队列", width=6.2)
    add_callout(doc, "重要边界",
                "语义检查只生成复核建议。管理员的“确认需治理”或“保留独立记录”只形成决定与操作记录，不会自动合并资产、删除来源或修改 Wiki。",
                fill=LIGHT_GOLD, accent=GOLD)

    # Architecture
    section_page(doc, "TECHNOLOGY", "技术能力被组织成可替换、可审计、可演进的产品架构",
                 "系统把外部模型能力放在同源服务网关之后，密钥不进入浏览器；正式知识由本地持久层、权限和审批服务控制。")
    add_two_col_table(doc, [
        ("交互层", "intelifar Web 工作空间：指挥台、文档、分析、资产、Wiki、Agent 与治理模块。"),
        ("业务层", "任务编排、权限校验、发布审批、版本、关系治理、分享、审计与备份。"),
        ("智能层", "MinerU 文档解析、DeepSeek 结构化分析与受控任务规划、Semantica 本地语义治理。"),
        ("知识层", "文档、IP 资产、原文证据、Wiki、版本和关系网络的统一对象模型。"),
        ("安全层", "同源网关、文件预检、限流、工作空间隔离、最小字段传递与审计留痕。"),
    ], header=("架构层", "当前实现"))
    add_heading(doc, "形成技术壁垒的四个组合", 2)
    for item in (
        "版面解析与专属结构提取结合：从“可读文本”走向“可治理 IP 对象”。",
        "结论与原文证据结合：让生成内容可核查，而不是只依赖模型表述。",
        "知识图谱与人工治理结合：关系能增长，但正式网络不被算法建议直接污染。",
        "自然语言任务与封闭能力边界结合：提升泛化效率，同时保留企业可控性。",
    ):
        add_bullet(doc, item)

    # Differentiation
    section_page(doc, "WHY INTELIFAR", "差异不在某一个模型，而在完整闭环和长期积累",
                 "基础模型和解析服务可以演进，但企业真正持续积累的是专属 Schema、证据链、资产关系、治理决定和使用反馈。")
    add_two_col_table(doc, [
        ("相对普通文档问答", "不仅回答问题，还形成可发布、可版本化、可追溯的资产和 Wiki。"),
        ("相对传统知识库", "从长文档自动提取结构、证据与关系，降低首次建库门槛。"),
        ("相对通用 Agent", "限定在文档 IP 与 Wiki 范围内，用受控工具和人工审批交付确定性结果。"),
        ("相对静态图谱", "关系来自真实资产与证据，支持权限过滤、待复核状态和持续治理。"),
    ], header=("比较维度", "intelifar 的产品选择"))
    add_callout(doc, "可积累的产品飞轮",
                "更多授权文档带来更完整的资产与证据；更多人工复核改善关系与治理质量；更好的知识网络提升搜索和任务交付价值；更高复用价值推动更多团队持续沉淀。",
                fill=LIGHT_TEAL, accent=TEAL)

    # Customer path
    section_page(doc, "CUSTOMER ADOPTION", "从一组真实资料开始，用可验收结果决定是否扩大",
                 "客户无需一次性迁移全部知识。更稳妥的方式是从一个高价值场景开始，以证据覆盖、复核效率和复用情况作为验收指标。")
    for step in (
        "选择样本：准备一组已授权、业务价值明确、包含版本与关系的真实资料。",
        "定义目标：选择资产盘点、技术尽调、项目知识复用、客户交付或风险核查之一。",
        "联合验证：检查提取结构、原文依据、关系网络、权限和任务交付是否满足要求。",
        "形成规范：确认 Schema、角色、密级、发布与分享流程。",
        "逐步扩展：增加文档量、团队和治理深度，并根据企业基础设施调整部署。",
    ):
        add_number(doc, step)
    add_heading(doc, "建议验收指标", 2)
    add_two_col_table(doc, [
        ("可信度", "关键结论是否能够回到正确原文位置；待核实项是否被明确标注。"),
        ("效率", "从接入文档到形成可复核资产与 Wiki 的时间和人工步骤。"),
        ("复用", "搜索、关系网络和任务助手是否帮助团队找到并使用已有知识。"),
        ("治理", "角色、审批、密级、分享、审计和备份是否符合当前组织要求。"),
    ], header=("维度", "验证问题"))

    # Investor lens
    section_page(doc, "INVESTOR VIEW", "产品扩展来自知识密度、组织协作和治理深度，而不只是模型调用量",
                 "intelifar 的价值可以随客户文档资产、使用岗位、知识关系和治理要求增长，并逐步形成难以迁移的组织知识网络。")
    add_two_col_table(doc, [
        ("文档规模", "从少量高价值资料扩展到跨项目、跨产品和跨年度知识。"),
        ("团队协作", "从负责人和内容维护者扩展到研发、法务、销售、交付与外部合作方。"),
        ("使用深度", "从解析与 Wiki 扩展到关系治理、任务代理、尽调、风险与生命周期管理。"),
        ("部署深度", "从小微企业单机持久工作空间扩展到企业身份、审计、存储和高可用适配。"),
    ], header=("增长轴", "产品扩展空间"))
    add_heading(doc, "投资判断所需的下一组证据", 2)
    for item in (
        "不同业务场景的样本文档到可复用资产的转化质量。",
        "知识编辑者、管理者和只读成员在真实工作中的持续使用频率。",
        "语义复核、关系确认和 Wiki 更新如何积累为更高质量的专属知识网络。",
        "从单团队试点向多团队推广时的实施成本、权限复杂度和商业转化路径。",
    ):
        add_bullet(doc, item)

    # Proof
    section_page(doc, "PROOF", "这不是概念图：核心链路已经通过真实服务和 UI 验证",
                 "截至 2026 年 8 月 12 日，当前版本已完成全量自动化测试、真实解析与模型链路、Semantica 本地语义检查以及多角色 UI E2E。")
    add_two_col_table(doc, [
        ("170 项测试", "服务、权限、发布、Wiki 版本、关系、Agent、备份、分享、语义复核与界面契约全部通过。"),
        ("50 个构建页面", "Astro 构建通过，桌面与移动端核心页面已形成。"),
        ("真实智能链路", "MinerU 与 DeepSeek 通过同源服务端链路运行，浏览器不持有 API 密钥。"),
        ("真实语义 E2E", "对 8 项真实工作空间资产完成检查，识别 4 条疑似重复候选并进入人工复核。"),
        ("安全检查", "高危依赖漏洞为 0，项目文本密钥泄漏扫描通过，真实 UI 控制台无错误。"),
    ], header=("验证证据", "已完成结果"))
    add_callout(doc, "生产部署说明",
                "小微企业版本已具备持久工作空间、权限、版本、审计、备份和安全分享基础。大型企业上线仍需结合现有统一身份、集中审计、外部恶意文件扫描、异地备份、网络与高可用体系进行实施。",
                fill=LIGHT_GOLD, accent=GOLD)

    # Close
    section_page(doc, "NEXT STEP", "带一组真实资料，验证它能否成为企业的下一批知识资产")
    add_para(doc, "一次有效的产品验证，不需要从“是否喜欢 AI”开始，而应从三个可回答的问题开始：",
             size=12, color=INK, after=16)
    add_callout(doc, "01  是否提取得准",
                "关键资产、关系与风险是否符合业务专家判断，并明确标出不确定项。")
    add_callout(doc, "02  是否说得清来源",
                "每项重要结论是否能够回到正确原文、版本与证据位置。", fill=LIGHT_TEAL, accent=TEAL)
    add_callout(doc, "03  是否管得住",
                "权限、审批、密级、分享、审计和任务边界是否符合企业责任链。", fill=LIGHT_GOLD, accent=GOLD)
    add_para(doc, "建议下一步：预约产品演示，或选择一组已授权资料开展联合验证。",
             size=14, color=PURPLE_DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             before=22, after=12)
    add_para(doc, "intelifar  |  让长文档成为可持续增长的企业知识资产",
             size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Metadata and accessibility basics
    props = doc.core_properties
    props.title = "intelifar 企业文档 IP 智能平台产品介绍"
    props.subject = "面向潜在客户与投资人的产品价值、功能、技术与合作说明"
    props.author = "intelifar"
    props.keywords = "intelifar, 企业文档, IP资产, Wiki, 知识图谱, Agent"
    props.comments = "基于 2026-08-12 已实现产品功能与验证证据"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
