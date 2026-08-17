from pathlib import Path
from zipfile import ZipFile
import re
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(r"C:\品味识别\intelifar-ip-wiki-graph")
DOCX = ROOT / "deliverables" / "intelifar-企业文档IP智能平台-产品介绍.docx"

doc = Document(DOCX)
assert len(doc.sections) == 1
section = doc.sections[0]
assert round(section.page_width.inches, 2) == 8.5
assert round(section.page_height.inches, 2) == 11.0
assert all(round(value.inches, 2) == 1.0 for value in (
    section.top_margin, section.right_margin, section.bottom_margin, section.left_margin
))

paragraph_text = "\n".join(p.text for p in doc.paragraphs)
assert "intelifar" in paragraph_text
assert "InteliFar" not in paragraph_text
assert "placeholder" not in paragraph_text.lower()
assert "TODO" not in paragraph_text
assert "codex-file-citation" not in paragraph_text
assert len(doc.inline_shapes) == 7, len(doc.inline_shapes)
assert len(doc.tables) >= 17, len(doc.tables)

for table in doc.tables:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    assert tbl_w is not None
    assert tbl_w.get(qn("w:type")) == "dxa"
    assert int(tbl_w.get(qn("w:w"))) == 9360
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == 120
    grid = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
    assert sum(grid) == 9360
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            assert tc_w is not None and tc_w.get(qn("w:type")) == "dxa"
            assert int(tc_w.get(qn("w:w"))) == grid[min(idx, len(grid) - 1)]

for shape in doc.inline_shapes:
    props = shape._inline.docPr
    assert props.get("descr"), "inline image missing alternative text"

for style_name, size in (("Normal", 11), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
    style = doc.styles[style_name]
    assert round(style.font.size.pt) == size

with ZipFile(DOCX) as zf:
    names = set(zf.namelist())
    assert "word/document.xml" in names
    assert "docProps/core.xml" in names
    doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert "INTELIFAR_SEMANTICA" not in doc_xml
    assert "apikey" not in doc_xml.lower()
    core = zf.read("docProps/core.xml").decode("utf-8")
    assert "intelifar" in core

print({
    "status": "PASS",
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "images": len(doc.inline_shapes),
    "pageGeometry": "Letter portrait, 1in margins",
    "tableGeometry": "9360 DXA fixed",
    "brand": "intelifar",
})
